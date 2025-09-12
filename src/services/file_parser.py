# File parsing utilities for different document formats
import io
import re
from typing import Union
import PyPDF2
from docx import Document
import logging

logger = logging.getLogger(__name__)

def parse_pdf(file_obj) -> str:
    """Extract text from PDF file"""
    try:
        # Reset file pointer
        file_obj.seek(0)
        
        # Try with PyPDF2
        pdf_reader = PyPDF2.PdfReader(file_obj)
        text = ""
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        # Clean up text
        text = clean_extracted_text(text)
        
        if not text.strip():
            logger.warning("No text extracted from PDF")
            return ""
        
        return text
        
    except Exception as e:
        logger.error(f"PDF parsing failed: {str(e)}")
        
        # Fallback: try alternative PDF parsing
        try:
            file_obj.seek(0)
            return parse_pdf_fallback(file_obj)
        except Exception as e2:
            logger.error(f"PDF fallback parsing also failed: {str(e2)}")
            return ""

def parse_pdf_fallback(file_obj) -> str:
    """Fallback PDF parsing using pdfminer"""
    try:
        from pdfminer.high_level import extract_text
        file_obj.seek(0)
        text = extract_text(file_obj)
        return clean_extracted_text(text)
    except ImportError:
        logger.warning("pdfminer not available for fallback parsing")
        return ""
    except Exception as e:
        logger.error(f"Fallback PDF parsing failed: {str(e)}")
        return ""

def parse_docx(file_obj) -> str:
    """Extract text from DOCX file"""
    try:
        file_obj.seek(0)
        doc = Document(file_obj)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        text = "\n".join(text_parts)
        return clean_extracted_text(text)
        
    except Exception as e:
        logger.error(f"DOCX parsing failed: {str(e)}")
        return ""

def parse_txt(file_obj) -> str:
    """Extract text from TXT file"""
    try:
        file_obj.seek(0)
        
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                file_obj.seek(0)
                content = file_obj.read()
                if isinstance(content, bytes):
                    text = content.decode(encoding)
                else:
                    text = content
                return clean_extracted_text(text)
            except UnicodeDecodeError:
                continue
        
        logger.error("Could not decode text file with any encoding")
        return ""
        
    except Exception as e:
        logger.error(f"TXT parsing failed: {str(e)}")
        return ""

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove page numbers and headers/footers patterns
    text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Page \d+', '', text, flags=re.IGNORECASE)
    
    # Fix common OCR/extraction issues
    text = re.sub(r'\s+([.,;:])', r'\1', text)  # Fix spaced punctuation
    text = re.sub(r'([.!?])\s*([A-Z])', r'\1\n\2', text)  # Add line breaks after sentences
    
    # Normalize bullet points
    text = re.sub(r'[•▪▫‣⁃]', '•', text)
    
    # Remove excessive line breaks but preserve structure
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()

def extract_text_by_filename(filename: str, file_obj) -> str:
    """Extract text based on file extension"""
    if not filename:
        return ""
    
    ext = filename.lower().split('.')[-1]
    
    if ext == 'pdf':
        return parse_pdf(file_obj)
    elif ext in ['docx', 'doc']:
        return parse_docx(file_obj)
    elif ext == 'txt':
        return parse_txt(file_obj)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

# Utility functions for text processing
def detect_sections_from_text(text: str) -> dict:
    """Advanced section detection with improved patterns"""
    
    sections = {}
    lines = text.split('\n')
    current_section = 'header'
    current_content = []
    
    # Enhanced section patterns
    section_patterns = {
        'contact': r'(contact|personal\s+info|reach\s+me)',
        'summary': r'(summary|profile|objective|about|overview)',
        'skills': r'(skills|technical|competencies|technologies|tools)',
        'experience': r'(experience|employment|work|career|professional)',
        'projects': r'(projects|portfolio|work\s+samples)',
        'education': r'(education|academic|qualifications|degrees?)',
        'certifications': r'(certifications?|licenses?|credentials|awards)',
        'volunteer': r'(volunteer|community|service)',
        'publications': r'(publications?|papers?|articles?)',
        'languages': r'(languages?|linguistic)',
    }
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        
        # Check if line is a section header
        section_found = None
        
        # Look for standalone headers (short lines, often capitalized)
        if len(line_clean) < 50 and (
            line_clean.isupper() or 
            line_clean.istitle() or 
            line_clean.endswith(':')
        ):
            for section, pattern in section_patterns.items():
                if re.search(pattern, line_clean.lower()):
                    section_found = section
                    break
        
        if section_found:
            # Save previous section
            if current_content:
                sections[current_section] = '\n'.join(current_content)
            current_section = section_found
            current_content = []
        else:
            current_content.append(line_clean)
    
    # Save last section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections

def extract_contact_info(text: str) -> dict:
    """Enhanced contact information extraction"""
    contact = {
        'name': '',
        'email': '',
        'phone': '',
        'location': '',
        'website': '',
        'linkedin': '',
        'github': ''
    }
    
    # Email extraction
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact['email'] = email_match.group()
    
    # Phone extraction (multiple formats)
    phone_patterns = [
        r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
        r'(\+\d{1,3}[-.\s]?)?(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})',
        r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            contact['phone'] = phone_match.group().strip()
            break
    
    # LinkedIn
    linkedin_patterns = [
        r'linkedin\.com/in/([A-Za-z0-9_-]+)',
        r'linkedin\.com/pub/([A-Za-z0-9_-]+)',
        r'www\.linkedin\.com/in/([A-Za-z0-9_-]+)'
    ]
    
    for pattern in linkedin_patterns:
        linkedin_match = re.search(pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"
            break
    
    # GitHub
    github_patterns = [
        r'github\.com/([A-Za-z0-9_-]+)',
        r'www\.github\.com/([A-Za-z0-9_-]+)'
    ]
    
    for pattern in github_patterns:
        github_match = re.search(pattern, text, re.IGNORECASE)
        if github_match:
            contact['github'] = f"github.com/{github_match.group(1)}"
            break
    
    # Website/Portfolio
    website_pattern = r'https?://(?:www\.)?([A-Za-z0-9.-]+\.[A-Za-z]{2,})'
    website_matches = re.findall(website_pattern, text, re.IGNORECASE)
    for match in website_matches:
        if 'linkedin' not in match.lower() and 'github' not in match.lower():
            contact['website'] = match
            break
    
    # Location (city, state patterns)
    location_patterns = [
        r'([A-Za-z\s]+,\s*[A-Z]{2}(?:\s+\d{5})?)',  # City, ST or City, ST ZIP
        r'([A-Za-z\s]+,\s*[A-Za-z\s]+)',  # City, Country
    ]
    
    for pattern in location_patterns:
        location_match = re.search(pattern, text)
        if location_match:
            location = location_match.group(1).strip()
            if len(location) < 50:  # Reasonable location length
                contact['location'] = location
                break
    
    # Name extraction (first few non-contact lines)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:  # Check first 5 lines
        # Skip lines with email, phone, or URLs
        if not re.search(r'@|phone|\d{3}|http|www|\.com', line.lower()):
            # Check if looks like a name (2-4 words, proper case)
            words = line.split()
            if 2 <= len(words) <= 4 and all(word.istitle() or word.isupper() for word in words):
                contact['name'] = line
                break
    
    return contact
