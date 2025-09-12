# Resume Generator Services
from dataclasses import dataclass
from typing import List, Dict, Optional, Tuple
import json
import re
from datetime import datetime
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import openai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pdfkit
from io import BytesIO

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Warning: spaCy model not found. Install with: python -m spacy download en_core_web_sm")
    nlp = None

@dataclass
class ResumeSchema:
    """Canonical resume data structure"""
    contact: Dict[str, str]
    summary: str
    skills: List[str]
    experience: List[Dict]
    projects: List[Dict]
    education: List[Dict]
    certifications: List[str]
    links: Dict[str, str]

class ResumeParser:
    """Extract and normalize resume content"""
    
    def __init__(self):
        self.section_patterns = {
            'contact': r'(contact|personal\s+info)',
            'summary': r'(summary|profile|objective|about)',
            'skills': r'(skills|technical|competencies|technologies)',
            'experience': r'(experience|employment|work|career)',
            'projects': r'(projects|portfolio)',
            'education': r'(education|academic|qualifications)',
            'certifications': r'(certifications?|licenses?|credentials)'
        }
        
    def parse_text(self, text: str) -> ResumeSchema:
        """Parse resume text into structured format"""
        sections = self._detect_sections(text)
        
        return ResumeSchema(
            contact=self._extract_contact(sections.get('contact', '')),
            summary=self._extract_summary(sections.get('summary', '')),
            skills=self._extract_skills(sections.get('skills', '')),
            experience=self._extract_experience(sections.get('experience', '')),
            projects=self._extract_projects(sections.get('projects', '')),
            education=self._extract_education(sections.get('education', '')),
            certifications=self._extract_certifications(sections.get('certifications', '')),
            links=self._extract_links(text)
        )
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """Split text into sections based on headers"""
        sections = {}
        lines = text.split('\n')
        current_section = 'contact'
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line is a section header
            section_found = None
            for section, pattern in self.section_patterns.items():
                if re.search(pattern, line.lower()):
                    section_found = section
                    break
            
            if section_found:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
            
        return sections
    
    def _extract_contact(self, text: str) -> Dict[str, str]:
        """Extract contact information"""
        contact = {'name': '', 'email': '', 'phone': '', 'location': ''}
        
        # Email
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone
        phone_match = re.search(r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        # Name (first non-email, non-phone line)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines:
            if not re.search(r'@|phone|\d{3}', line.lower()) and len(line.split()) <= 4:
                contact['name'] = line
                break
                
        return contact
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills"""
        # Common technical skills patterns
        tech_patterns = [
            r'\b(Python|Java|JavaScript|TypeScript|React|Angular|Vue|Node\.js|Django|Flask|Spring|Kubernetes|Docker|AWS|Azure|GCP)\b',
            r'\b(SQL|PostgreSQL|MySQL|MongoDB|Redis|GraphQL|REST|API)\b',
            r'\b(Git|Jenkins|CI/CD|DevOps|Linux|Windows|MacOS)\b'
        ]
        
        skills = set()
        for pattern in tech_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            skills.update(matches)
        
        # Also split on common delimiters
        delimited_skills = re.split(r'[,•\n\t|]+', text)
        for skill in delimited_skills:
            skill = skill.strip()
            if skill and len(skill) < 30:  # Reasonable skill length
                skills.add(skill)
        
        return list(skills)[:20]  # Limit to top 20 skills
    
    def _extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience"""
        experiences = []
        # This is a simplified version - would need more sophisticated parsing
        sections = text.split('\n\n')
        
        for section in sections:
            if len(section.split('\n')) >= 2:
                lines = section.split('\n')
                exp = {
                    'company': '',
                    'role': '',
                    'start': '',
                    'end': '',
                    'bullets': []
                }
                
                # Try to parse first line as role/company
                if lines:
                    first_line = lines[0].strip()
                    if '|' in first_line or '-' in first_line:
                        parts = re.split(r'[|-]', first_line, 1)
                        exp['role'] = parts[0].strip()
                        if len(parts) > 1:
                            exp['company'] = parts[1].strip()
                    else:
                        exp['role'] = first_line
                
                # Extract bullets
                for line in lines[1:]:
                    line = line.strip()
                    if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        exp['bullets'].append(line[1:].strip())
                    elif line and not re.search(r'\d{4}', line):  # Not a date line
                        exp['bullets'].append(line)
                
                if exp['role'] or exp['bullets']:
                    experiences.append(exp)
        
        return experiences
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary"""
        # Clean up the text and return first paragraph
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            return ' '.join(lines)
        return ""
    
    def _extract_projects(self, text: str) -> List[Dict]:
        """Extract projects"""
        # Similar to experience parsing but for projects
        return []
    
    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education"""
        return []
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        return []
    
    def _extract_links(self, text: str) -> Dict[str, str]:
        """Extract social/professional links"""
        links = {}
        
        # GitHub
        github_match = re.search(r'github\.com/([A-Za-z0-9_-]+)', text)
        if github_match:
            links['github'] = f"github.com/{github_match.group(1)}"
        
        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/([A-Za-z0-9_-]+)', text)
        if linkedin_match:
            links['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"
        
        return links

class JobDescriptionAnalyzer:
    """Extract keywords and requirements from job descriptions"""
    
    def __init__(self):
        self.skill_taxonomy = {
            'languages': ['Python', 'Java', 'JavaScript', 'TypeScript', 'Go', 'Rust', 'C++', 'C#'],
            'frameworks': ['React', 'Angular', 'Vue', 'Django', 'Flask', 'Spring', 'Express'],
            'databases': ['PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Elasticsearch'],
            'cloud': ['AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes'],
            'tools': ['Git', 'Jenkins', 'Terraform', 'Ansible']
        }
    
    def extract_keywords(self, jd_text: str, use_llm: bool = False) -> List[Dict]:
        """Extract keywords from job description"""
        if use_llm:
            return self._extract_keywords_llm(jd_text)
        else:
            return self._extract_keywords_heuristic(jd_text)
    
    def _extract_keywords_heuristic(self, text: str) -> List[Dict]:
        """Extract keywords using TF-IDF and skill taxonomy"""
        keywords = []
        
        # Check against skill taxonomy
        text_lower = text.lower()
        for category, skills in self.skill_taxonomy.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    keywords.append({
                        'term': skill,
                        'category': category,
                        'importance': self._calculate_importance(skill, text)
                    })
        
        # Use spaCy for additional noun phrases
        if nlp:
            doc = nlp(text)
            for chunk in doc.noun_chunks:
                if len(chunk.text.split()) <= 3:  # Short phrases only
                    keywords.append({
                        'term': chunk.text,
                        'category': 'general',
                        'importance': 0.5
                    })
        
        # Remove duplicates and sort by importance
        seen = set()
        unique_keywords = []
        for kw in sorted(keywords, key=lambda x: x['importance'], reverse=True):
            if kw['term'].lower() not in seen:
                seen.add(kw['term'].lower())
                unique_keywords.append(kw)
        
        return unique_keywords[:30]  # Top 30 keywords
    
    def _extract_keywords_llm(self, text: str) -> List[Dict]:
        """Extract keywords using LLM"""
        prompt = """Extract must-have skills, tools, frameworks, and technologies from this job description.
        Return a JSON array with objects containing: term, category (skill/tool/framework/language/methodology), importance (0-1).
        
        Job Description:
        {text}
        
        Response (JSON only):"""
        
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt.format(text=text)}],
                max_tokens=500,
                temperature=0.1
            )
            
            result = json.loads(response.choices[0].message.content)
            return result if isinstance(result, list) else []
        except Exception as e:
            print(f"LLM extraction failed: {e}")
            return self._extract_keywords_heuristic(text)
    
    def _calculate_importance(self, term: str, text: str) -> float:
        """Calculate importance score for a term"""
        count = text.lower().count(term.lower())
        return min(count * 0.2, 1.0)

class BulletRewriter:
    """Rewrite bullet points for maximum impact"""
    
    def __init__(self):
        self.action_verbs = [
            'Developed', 'Built', 'Created', 'Designed', 'Implemented', 'Led', 'Managed',
            'Optimized', 'Improved', 'Reduced', 'Increased', 'Delivered', 'Achieved'
        ]
        
    def rewrite_bullets(self, bullets: List[str], jd_text: str = "", tone: str = "professional") -> List[Dict]:
        """Rewrite bullets for impact and ATS optimization"""
        rewrites = []
        
        for bullet in bullets:
            original = bullet.strip()
            if not original:
                continue
                
            rewritten = self._rewrite_single_bullet(original, jd_text, tone)
            improvements = self._identify_improvements(original, rewritten)
            
            rewrites.append({
                'original': original,
                'rewritten': rewritten,
                'improvements': improvements
            })
        
        return rewrites
    
    def _rewrite_single_bullet(self, bullet: str, jd_text: str, tone: str) -> str:
        """Rewrite a single bullet point"""
        # Remove first-person pronouns
        bullet = re.sub(r'\b(I|my|me)\b', '', bullet, flags=re.IGNORECASE).strip()
        
        # Ensure starts with action verb
        if not any(bullet.lower().startswith(verb.lower()) for verb in self.action_verbs):
            # Try to find existing verb and move it to front
            for verb in self.action_verbs:
                if verb.lower() in bullet.lower():
                    bullet = f"{verb} {bullet.replace(verb, '', 1)}"
                    break
            else:
                # Default action verb
                bullet = f"Developed {bullet}"
        
        # Clean up extra spaces
        bullet = re.sub(r'\s+', ' ', bullet).strip()
        
        # Ensure proper sentence structure
        if not bullet.endswith('.'):
            bullet += '.'
        
        return bullet
    
    def _identify_improvements(self, original: str, rewritten: str) -> List[str]:
        """Identify what was improved"""
        improvements = []
        
        if 'I ' in original and 'I ' not in rewritten:
            improvements.append("Removed first-person pronouns")
        
        if len(rewritten.split()) < len(original.split()):
            improvements.append("Made more concise")
        
        if any(verb in rewritten for verb in self.action_verbs):
            improvements.append("Started with strong action verb")
        
        return improvements

class ATSValidator:
    """Validate resume for ATS compliance"""
    
    def validate(self, resume_data: ResumeSchema) -> List[Dict]:
        """Run ATS compliance checks"""
        issues = []
        
        # Check contact information
        if not resume_data.contact.get('email'):
            issues.append({
                'severity': 'critical',
                'issue': 'Missing email address',
                'fix': 'Add professional email address'
            })
        
        if not resume_data.contact.get('phone'):
            issues.append({
                'severity': 'high',
                'issue': 'Missing phone number',
                'fix': 'Add phone number with area code'
            })
        
        # Check sections
        if not resume_data.experience:
            issues.append({
                'severity': 'critical',
                'issue': 'No work experience found',
                'fix': 'Add work experience section'
            })
        
        if len(resume_data.skills) < 5:
            issues.append({
                'severity': 'medium',
                'issue': 'Few technical skills listed',
                'fix': 'Add more relevant technical skills'
            })
        
        # Check bullet points
        for exp in resume_data.experience:
            if len(exp.get('bullets', [])) < 2:
                issues.append({
                    'severity': 'medium',
                    'issue': f'Few accomplishments for {exp.get("role", "position")}',
                    'fix': 'Add 3-5 bullet points per role'
                })
        
        return issues

class ResumeGenerator:
    """Generate ATS-optimized resumes in multiple formats"""
    
    def __init__(self):
        self.templates = {
            'plain': self._generate_plain_template,
            'compact': self._generate_compact_template,
            'engineer': self._generate_engineer_template
        }
    
    def generate(self, resume_data: ResumeSchema, template: str = 'plain', format_type: str = 'html') -> bytes:
        """Generate resume in specified format"""
        if template not in self.templates:
            raise ValueError(f"Unknown template: {template}")
        
        html_content = self.templates[template](resume_data)
        
        if format_type == 'html':
            return html_content.encode('utf-8')
        elif format_type == 'pdf':
            return self._html_to_pdf(html_content)
        elif format_type == 'docx':
            return self._generate_docx(resume_data, template)
        else:
            raise ValueError(f"Unknown format: {format_type}")
    
    def _generate_plain_template(self, data: ResumeSchema) -> str:
        """Generate plain ATS template"""
        return f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ 
            font-family: 'Times New Roman', serif; 
            font-size: 11pt; 
            margin: 0.75in; 
            line-height: 1.2; 
            color: #000;
        }}
        .header {{ text-align: center; margin-bottom: 20pt; }}
        .name {{ font-size: 16pt; font-weight: bold; margin-bottom: 4pt; }}
        .contact {{ font-size: 10pt; }}
        .section {{ margin-bottom: 12pt; }}
        .section-title {{ 
            font-size: 12pt; 
            font-weight: bold; 
            border-bottom: 1px solid #000; 
            margin-bottom: 6pt; 
            text-transform: uppercase;
        }}
        .job {{ margin-bottom: 8pt; }}
        .job-header {{ font-weight: bold; }}
        .job-details {{ font-style: italic; margin-bottom: 2pt; }}
        ul {{ margin: 2pt 0; padding-left: 20pt; }}
        li {{ margin-bottom: 1pt; }}
        p {{ margin: 0 0 4pt 0; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="name">{data.contact.get('name', '')}</div>
        <div class="contact">
            {data.contact.get('email', '')} • {data.contact.get('phone', '')} • {data.contact.get('location', '')}
        </div>
        <div class="contact">
            {' • '.join([v for v in data.links.values() if v])}
        </div>
    </div>

    {f'''
    <div class="section">
        <div class="section-title">SUMMARY</div>
        <p>{data.summary}</p>
    </div>
    ''' if data.summary else ''}

    {f'''
    <div class="section">
        <div class="section-title">TECHNICAL SKILLS</div>
        <p>{', '.join(data.skills)}</p>
    </div>
    ''' if data.skills else ''}

    {f'''
    <div class="section">
        <div class="section-title">EXPERIENCE</div>
        {''.join([f"""
        <div class="job">
            <div class="job-header">{exp.get('role', '')}</div>
            <div class="job-details">{exp.get('company', '')} • {exp.get('start', '')} - {exp.get('end', '')}</div>
            <ul>
                {''.join([f"<li>{bullet}</li>" for bullet in exp.get('bullets', [])])}
            </ul>
        </div>
        """ for exp in data.experience])}
    </div>
    ''' if data.experience else ''}

    {f'''
    <div class="section">
        <div class="section-title">EDUCATION</div>
        {''.join([f"<p><strong>{edu.get('school', '')}</strong> • {edu.get('degree', '')} • {edu.get('grad', '')}</p>" for edu in data.education])}
    </div>
    ''' if data.education else ''}

    {f'''
    <div class="section">
        <div class="section-title">CERTIFICATIONS</div>
        <p>{', '.join(data.certifications)}</p>
    </div>
    ''' if data.certifications else ''}
</body>
</html>
        """.strip()
    
    def _generate_compact_template(self, data: ResumeSchema) -> str:
        """Generate compact template for space optimization"""
        # Similar to plain but with tighter spacing
        return self._generate_plain_template(data).replace('margin-bottom: 12pt', 'margin-bottom: 8pt')
    
    def _generate_engineer_template(self, data: ResumeSchema) -> str:
        """Generate engineer-focused template with skills first"""
        # Reorder sections to put skills first
        template = self._generate_plain_template(data)
        # This would involve more complex template manipulation
        return template
    
    def _html_to_pdf(self, html_content: str) -> bytes:
        """Convert HTML to PDF"""
        try:
            pdf_buffer = BytesIO()
            # Using pdfkit (requires wkhtmltopdf)
            options = {
                'page-size': 'Letter',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None
            }
            
            pdf_data = pdfkit.from_string(html_content, False, options=options)
            return pdf_data
        except Exception as e:
            print(f"PDF generation failed: {e}")
            return b""
    
    def _generate_docx(self, data: ResumeSchema, template: str) -> bytes:
        """Generate DOCX file"""
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)
        
        # Create header style
        header_style = styles.add_style('Header', WD_STYLE_TYPE.PARAGRAPH)
        header_style.base_style = normal_style
        header_style.font.size = Pt(16)
        header_style.font.bold = True
        header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create section title style
        section_style = styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
        section_style.base_style = normal_style
        section_style.font.size = Pt(12)
        section_style.font.bold = True
        
        # Add content
        # Name
        name_para = doc.add_paragraph(data.contact.get('name', ''), style='Header')
        
        # Contact info
        contact_info = f"{data.contact.get('email', '')} • {data.contact.get('phone', '')} • {data.contact.get('location', '')}"
        contact_para = doc.add_paragraph(contact_info)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Links
        if data.links:
            links_text = ' • '.join([v for v in data.links.values() if v])
            links_para = doc.add_paragraph(links_text)
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if data.summary:
            doc.add_paragraph('SUMMARY', style='SectionTitle')
            doc.add_paragraph(data.summary)
        
        # Skills
        if data.skills:
            doc.add_paragraph('TECHNICAL SKILLS', style='SectionTitle')
            doc.add_paragraph(', '.join(data.skills))
        
        # Experience
        if data.experience:
            doc.add_paragraph('EXPERIENCE', style='SectionTitle')
            for exp in data.experience:
                # Role
                role_para = doc.add_paragraph()
                role_run = role_para.add_run(exp.get('role', ''))
                role_run.bold = True
                
                # Company and dates
                details_para = doc.add_paragraph()
                details_run = details_para.add_run(f"{exp.get('company', '')} • {exp.get('start', '')} - {exp.get('end', '')}")
                details_run.italic = True
                
                # Bullets
                for bullet in exp.get('bullets', []):
                    bullet_para = doc.add_paragraph(bullet, style='List Bullet')
        
        # Save to bytes
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer.getvalue()

# LLM Prompts
LLM_PROMPTS = {
    'extract_keywords': """Extract must-have skills, tools, frameworks, and technologies from this job description.
Focus on technical requirements and hard skills. Return a JSON array with objects containing:
- term: the skill/technology name
- category: one of [skill, tool, framework, language, methodology, certification]
- importance: score from 0-1 based on how critical it seems

Job Description:
{job_description}

Response (JSON only):""",

    'rewrite_bullet': """Rewrite this bullet point to follow the pattern: Action Verb + Task + Quantified Impact.
Keep it under 24 words. Preserve technical terms. Remove first-person pronouns.

Original: {bullet}
Job context: {job_description}

Rewritten bullet:""",

    'generate_summary': """Write a professional summary (3 lines max) emphasizing the top 5 skills from this job description.
No first-person pronouns. Include years of experience if mentioned.

Resume data: {resume_data}
Job requirements: {job_description}

Summary:"""
}
